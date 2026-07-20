import json
import uuid
from datetime import datetime
from pathlib import Path

from flask import Blueprint, current_app, jsonify, render_template, request, send_file
from werkzeug.utils import secure_filename

from database import get_db
from file_parser import parse_file
from file_validation import validate_cv_file
from job_scraper import parse_job_text, scrape_linkedin_job, validate_linkedin_url
from repositories import JobSearchRepository
from services import JobSearchRepository as JobSearchDeleteRepository
from routes.helpers import allowed_file, extractive_cv_gen, nlp_engine, remove_upload_file, validate_saved_cv

job_search_bp = Blueprint("job_search", __name__)

@job_search_bp.route('/api/job-search/delete-session/<session_id>', methods=['DELETE'])
def job_search_delete_session(session_id):

    repository = JobSearchRepository(get_db())
    JobSearchDeleteRepository(repository._db).delete_session(session_id)
    repository.commit()
    return jsonify({'success': True, 'message': 'Oturum silindi.'})


@job_search_bp.route('/job-search')
def job_search():

    repository = JobSearchRepository(get_db())

    profiles = repository.execute(
        "SELECT id, original_filename, extracted_skills, created_at FROM user_profiles ORDER BY created_at DESC"
    ).fetchall()

    sessions = repository.execute("""
        SELECT s.*, p.original_filename
        FROM job_search_sessions s
        JOIN user_profiles p ON s.profile_id = p.id
        ORDER BY s.created_at DESC LIMIT 10
    """).fetchall()
    return render_template('job_search.html', profiles=profiles, sessions=sessions)

@job_search_bp.route('/api/job-search/upload-cv', methods=['POST'])
def job_search_upload_cv():

    filepath = None

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'Dosya seçilmedi.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'Dosya seçilmedi.'}), 400

    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'Sadece PDF ve DOCX dosyaları desteklenmektedir.'}), 400

    try:

        filename = secure_filename(file.filename)
        unique_name = f"{uuid.uuid4().hex[:8]}_{filename}"
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], unique_name)
        file.save(filepath)

        if not validate_saved_cv(Path(filepath)):
            return jsonify({'success': False, 'error': 'Dosya içeriği geçersiz veya desteklenmiyor.'}), 400

        parse_result = parse_file(filepath)
        raw_text = parse_result['text']
        warnings = parse_result['warnings']
        if not raw_text or len(raw_text.strip()) < 50:
            remove_upload_file(Path(filepath))
            return jsonify({'success': False, 'error': 'PDF/DOCX\'den yeterli metin çıkarılamadı.'}), 400

        profile_data = extractive_cv_gen().parse_cv(raw_text)

        if not profile_data:
            remove_upload_file(Path(filepath))
            return jsonify({'success': False, 'error': 'Profil parse edilemedi.'}), 500

        skills = nlp_engine().extract_skills(raw_text)

        profile_id = str(uuid.uuid4())
        repository = JobSearchRepository(get_db())
        repository.execute(
            """INSERT INTO user_profiles (id, original_filename, original_text,
               profile_data, extracted_skills, created_at)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (profile_id, filename, raw_text,
             json.dumps(profile_data, ensure_ascii=False),
             json.dumps(skills),
             datetime.now().isoformat())
        )
        repository.commit()

        remove_upload_file(Path(filepath))

        return {
            'success': True,
            'profile_id': profile_id,
            'profile': profile_data,
            'skills': skills,
            'filename': filename,
            'warnings': warnings
        }

    except Exception:
        current_app.logger.exception('Job search CV upload failed')
        remove_upload_file(Path(filepath) if filepath else None)
        return jsonify({'success': False, 'error': 'CV işlenemedi. Lütfen tekrar deneyin.'}), 500

@job_search_bp.route('/api/job-search/parse-job', methods=['POST'])
def job_search_parse_job():

    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'error': 'Veri gönderilmedi.'}), 400

    url = data.get('url', '').strip()
    manual_text = data.get('manual_text', '').strip()
    manual_title = data.get('manual_title', '').strip()
    manual_company = data.get('manual_company', '').strip()

    if manual_text:

        result = parse_job_text(manual_text, manual_title, manual_company)
    elif url:

        if not validate_linkedin_url(url):
            return jsonify({
                'success': False,
                'error': 'Geçersiz LinkedIn URL formatı. Lütfen geçerli bir LinkedIn iş ilanı linki girin.'
            }), 400
        result = scrape_linkedin_job(url)
    else:
        return jsonify({'success': False, 'error': 'LinkedIn URL veya ilan metni girilmelidir.'}), 400

    if result['success'] and result.get('description'):
        extracted_skills = nlp_engine().extract_skills(result['description'])
        result['extracted_skills'] = extracted_skills

    return jsonify(result)

@job_search_bp.route('/api/job-search/generate-cv', methods=['POST'])
def job_search_generate_cv():

    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'error': 'Veri gönderilmedi.'}), 400

    profile_id = data.get('profile_id', '').strip()
    job_data = data.get('job_data')

    if not profile_id:
        return jsonify({'success': False, 'error': 'Profil ID gerekli.'}), 400
    if not job_data:
        return jsonify({'success': False, 'error': 'İlan bilgileri gerekli.'}), 400

    repository = JobSearchRepository(get_db())
    profile = repository.execute("SELECT * FROM user_profiles WHERE id = ?", (profile_id,)).fetchone()
    if not profile:
        return jsonify({'success': False, 'error': 'Profil bulunamadı.'}), 404

    profile_data = json.loads(profile['profile_data'])

    try:
        optimized_cv_data = extractive_cv_gen().generate_tailored_cv(profile_data, job_data)
    except Exception:
        current_app.logger.exception('CV generation failed')
        return jsonify({'success': False, 'error': 'CV oluşturulamadı. Lütfen tekrar deneyin.'}), 500

    session_id = str(uuid.uuid4())
    repository.execute(
        """INSERT INTO job_search_sessions (id, profile_id, job_url, job_data,
           optimized_cv, status, created_at)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (session_id, profile_id,
         job_data.get('url', ''),
         json.dumps(job_data, ensure_ascii=False),
         json.dumps(optimized_cv_data, ensure_ascii=False),
         'completed',
         datetime.now().isoformat())
    )
    repository.commit()

    return jsonify({
        'success': True,
        'session_id': session_id,
        'cv_data': optimized_cv_data
    })

def generate_docx_buffer(cv_data):

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv_data.get('full_name', 'İsimsiz').upper())
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact = cv_data.get('contact', {})
    contact_parts = []
    if contact.get('email'): contact_parts.append(contact['email'])
    if contact.get('phone'): contact_parts.append(contact['phone'])
    if contact.get('location'): contact_parts.append(contact['location'])
    if contact.get('linkedin'): contact_parts.append(contact['linkedin'])

    if contact_parts:
        contact_p = doc.add_paragraph(' • '.join(contact_parts))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = cv_data.get('headers', {})
    summary = cv_data.get('professional_summary', '')
    if summary:
        doc.add_heading(headers.get('summary', 'SUMMARY'), level=1)
        doc.add_paragraph(summary)

    experiences = cv_data.get('experience', [])
    if experiences:
        doc.add_heading(headers.get('experience', 'EXPERIENCE'), level=1)
        for exp in experiences:
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp.get('title', '')}")
            title_run.bold = True
            p.add_run(f" — {exp.get('company', '')}")

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(2)
            meta_text = []
            if exp.get('location'): meta_text.append(exp['location'])
            if exp.get('period'): meta_text.append(exp['period'])
            meta_p.add_run(' | '.join(meta_text)).italic = True

            for ach in exp.get('achievements', []):
                bullet = doc.add_paragraph(ach, style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.25)

    education = cv_data.get('education', [])
    if education:
        doc.add_heading(headers.get('education', 'EDUCATION'), level=1)
        for edu in education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')}").bold = True
            p.add_run(f" | {edu.get('institution', '')}")
            if edu.get('period'):
                p.add_run(f" ({edu.get('period', '')})").italic = True
            if edu.get('details'):
                doc.add_paragraph(edu['details'])

    skills = cv_data.get('skills', {})
    if skills:
        doc.add_heading(headers.get('skills', 'SKILLS'), level=1)
        skill_groups = skills.get('groups', [])
        if skill_groups:
            for group in skill_groups:
                doc.add_paragraph(", ".join(group))
        else:
            if skills.get('primary'):
                p = doc.add_paragraph()
                p.add_run("Temel: ").bold = True
                p.add_run(", ".join(skills['primary']))
            if skills.get('secondary'):
                p = doc.add_paragraph()
                p.add_run("Ek: ").bold = True
                p.add_run(", ".join(skills['secondary']))
            if skills.get('tools'):
                p = doc.add_paragraph()
                p.add_run("Araçlar: ").bold = True
                p.add_run(", ".join(skills['tools']))

    if cv_data.get('languages'):
        doc.add_heading(headers.get('languages', 'LANGUAGES'), level=1)
        langs = [f"{l.get('language','')} ({l.get('level','')})" for l in cv_data['languages']]
        p = doc.add_paragraph()
        p.add_run(", ".join(langs))

    if cv_data.get('certifications'):
        doc.add_heading(headers.get('certifications', 'CERTIFICATIONS'), level=1)
        for cert in cv_data.get('certifications', []):
            cert_text = f"{cert.get('name','')} - {cert.get('issuer','')}"
            if cert.get('year'): cert_text += f" ({cert['year']})"
            doc.add_paragraph(cert_text, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@job_search_bp.route('/api/job-search/download-cv/<session_id>')
def job_search_download_cv(session_id):

    file_format = request.args.get('format', 'pdf').lower()

    repository = JobSearchRepository(get_db())
    session = repository.execute("SELECT * FROM job_search_sessions WHERE id = ?", (session_id,)).fetchone()
    if not session:
        return jsonify({'success': False, 'error': 'Oturum bulunamadı.'}), 404

    cv_data = json.loads(session['optimized_cv'])
    name = cv_data.get('full_name', 'cv_ats_optimized')
    safe_name = name.replace(' ', '_').lower()[:30]

    if file_format == 'docx':
        buffer = generate_docx_buffer(cv_data)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_name}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    repository = JobSearchRepository(get_db())
    session = repository.execute("SELECT * FROM job_search_sessions WHERE id = ?", (session_id,)).fetchone()
    if not session:
        return jsonify({'success': False, 'error': 'Oturum bulunamadı.'}), 404

    cv_data = json.loads(session['optimized_cv'])

    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=15*mm, bottomMargin=15*mm
    )

    font_name = 'Helvetica'
    try:
        import os
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        arial_path = r'C:\Windows\Fonts\arial.ttf'
        arial_bold_path = r'C:\Windows\Fonts\arialbd.ttf'
        arial_italic_path = r'C:\Windows\Fonts\ariali.ttf'

        if os.path.exists(arial_path):
            pdfmetrics.registerFont(TTFont('Arial', arial_path))
            pdfmetrics.registerFont(TTFont('Arial-Bold', arial_bold_path))
            pdfmetrics.registerFont(TTFont('Arial-Italic', arial_italic_path))
            font_name = 'Arial'
    except:
        pass

    font_bold = f"{font_name}-Bold" if font_name == 'Arial' else 'Helvetica-Bold'
    font_italic = f"{font_name}-Italic" if font_name == 'Arial' else 'Helvetica-Oblique'

    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name='CVName', fontSize=18, fontName=font_bold,
        spaceAfter=6, textColor=HexColor('#000000'), alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        name='CVContact', fontSize=10, fontName=font_name,
        spaceAfter=10, textColor=HexColor('#333333'), alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        name='CVSection', fontSize=12, fontName=font_bold,
        spaceBefore=12, spaceAfter=4, textColor=HexColor('#000000'),
        borderPadding=(0, 0, 2, 0)
    ))
    styles.add(ParagraphStyle(
        name='CVSubtitle', fontSize=11, fontName=font_bold,
        spaceAfter=2, textColor=HexColor('#000000')
    ))
    styles.add(ParagraphStyle(
        name='CVMeta', fontSize=9, fontName=font_italic,
        spaceAfter=3, textColor=HexColor('#444444')
    ))
    styles.add(ParagraphStyle(
        name='CVBody', fontSize=10.5, fontName=font_name,
        spaceAfter=4, textColor=HexColor('#000000'), leading=14,
        alignment=TA_LEFT
    ))
    styles.add(ParagraphStyle(
        name='CVBullet', fontSize=10.5, fontName=font_name,
        spaceAfter=3, textColor=HexColor('#000000'), leading=13,
        leftIndent=15, bulletIndent=5
    ))

    elements = []

    name = cv_data.get('full_name', 'İsimsiz')
    elements.append(Paragraph(name, styles['CVName']))

    contact = cv_data.get('contact', {})
    contact_parts = []
    if contact.get('email'):
        contact_parts.append(contact['email'])
    if contact.get('phone'):
        contact_parts.append(contact['phone'])
    if contact.get('location'):
        contact_parts.append(contact['location'])
    if contact.get('linkedin'):
        contact_parts.append(contact['linkedin'])
    if contact_parts:
        elements.append(Paragraph(' • '.join(contact_parts), styles['CVContact']))

    elements.append(HRFlowable(width='100%', thickness=1, color=HexColor('#c0c0e0'), spaceAfter=8))

    headers = cv_data.get('headers', {})
    summary = cv_data.get('professional_summary', '')
    if summary:
        elements.append(Paragraph(headers.get('summary', 'PROFESSIONAL SUMMARY'), styles['CVSection']))
        elements.append(Paragraph(summary, styles['CVBody']))

    experiences = cv_data.get('experience', [])
    if experiences:
        elements.append(Paragraph(headers.get('experience', 'EXPERIENCE'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for exp in experiences:
            title_company = f"<b>{exp.get('title', '')}</b> — {exp.get('company', '')}"
            elements.append(Paragraph(title_company, styles['CVSubtitle']))
            meta_parts = []
            if exp.get('location'):
                meta_parts.append(exp['location'])
            if exp.get('period'):
                meta_parts.append(exp['period'])
            if meta_parts:
                elements.append(Paragraph(' | '.join(meta_parts), styles['CVMeta']))
            for ach in exp.get('achievements', []):
                elements.append(Paragraph(f'• {ach}', styles['CVBullet']))
            elements.append(Spacer(1, 4))

    education = cv_data.get('education', [])
    if education:
        elements.append(Paragraph(headers.get('education', 'EDUCATION'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for edu in education:
            elements.append(Paragraph(f"<b>{edu.get('degree', '')}</b>", styles['CVSubtitle']))
            meta_parts = []
            if edu.get('institution'):
                meta_parts.append(edu['institution'])
            if edu.get('period'):
                meta_parts.append(edu['period'])
            if meta_parts:
                elements.append(Paragraph(' | '.join(meta_parts), styles['CVMeta']))
            if edu.get('details'):
                elements.append(Paragraph(edu['details'], styles['CVBody']))
            elements.append(Spacer(1, 3))

    skills = cv_data.get('skills', {})
    if skills:
        elements.append(Paragraph(headers.get('skills', 'SKILLS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))

        skill_groups = skills.get('groups', [])
        if skill_groups:
            for group in skill_groups:
                elements.append(Paragraph(', '.join(group), styles['CVBody']))
        else:

            if skills.get('primary'):
                elements.append(Paragraph(f"<b>Temel:</b> {', '.join(skills['primary'])}", styles['CVBody']))
            if skills.get('secondary'):
                elements.append(Paragraph(f"<b>Ek:</b> {', '.join(skills['secondary'])}", styles['CVBody']))
            if skills.get('tools'):
                elements.append(Paragraph(f"<b>Araçlar:</b> {', '.join(skills['tools'])}", styles['CVBody']))

    languages = cv_data.get('languages', [])
    if languages:
        elements.append(Paragraph(headers.get('languages', 'LANGUAGES'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        lang_text = ', '.join([f"{l.get('language', '')} ({l.get('level', '')})" for l in languages])
        elements.append(Paragraph(lang_text, styles['CVBody']))

    certs = cv_data.get('certifications', [])
    if certs:
        elements.append(Paragraph(headers.get('certifications', 'CERTIFICATIONS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for cert in certs:
            cert_text = cert.get('name', '')
            if cert.get('issuer'):
                cert_text += f" — {cert['issuer']}"
            if cert.get('year'):
                cert_text += f" ({cert['year']})"
            elements.append(Paragraph(f'• {cert_text}', styles['CVBullet']))

    projects = cv_data.get('projects', [])
    if projects:
        elements.append(Paragraph(headers.get('projects', 'PROJECTS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for proj in projects:
            elements.append(Paragraph(f"<b>{proj.get('name', '')}</b>", styles['CVSubtitle']))
            if proj.get('description'):
                elements.append(Paragraph(proj['description'], styles['CVBody']))
            if proj.get('technologies'):
                elements.append(Paragraph(f"Teknolojiler: {', '.join(proj['technologies'])}", styles['CVMeta']))
            elements.append(Spacer(1, 3))

    doc.build(elements)
    buffer.seek(0)

    safe_name = name.replace(' ', '_').lower()[:30]
    pdf_filename = f"cv_{safe_name}_{session_id[:8]}.pdf"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=pdf_filename,
        mimetype='application/pdf'
    )

@job_search_bp.route('/api/job-search/delete-profile/<profile_id>', methods=['DELETE'])
def delete_profile(profile_id):

    repository = JobSearchRepository(get_db())
    JobSearchDeleteRepository(repository._db).delete_profile(profile_id)
    repository.commit()
    return jsonify({'success': True, 'message': 'Profil silindi.'})

