import os
import json
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path

import numpy as np
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from werkzeug.utils import secure_filename

from config import Config
from database import get_db, init_app, init_db
from file_parser import parse_file
from file_validation import validate_cv_file
from nlp_engine import NLPEngine, clean_text
from job_scraper import scrape_linkedin_job, parse_job_text, validate_linkedin_url
from extractive_cv import ExtractiveCVGenerator

Config.validate()
app = Flask(__name__)
app.config.from_object(Config)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

nlp_engine = NLPEngine(
    sbert_model_name=app.config['SBERT_MODEL'],
    spacy_model_name=app.config['SPACY_MODEL']
)

extractive_cv_gen = ExtractiveCVGenerator(nlp_engine=nlp_engine)
init_app(app)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def get_upload_path(stored_filename):
    """Return a file path only when it points to a direct child of uploads."""
    if not stored_filename:
        return None

    upload_dir = Path(app.config['UPLOAD_FOLDER']).resolve()
    candidate = (upload_dir / stored_filename).resolve()
    if candidate.parent != upload_dir:
        return None
    return candidate


def remove_upload_file(filepath):
    """Delete a temporary upload without exposing filesystem errors to users."""
    try:
        if filepath and filepath.is_file():
            filepath.unlink()
    except OSError:
        app.logger.exception('Unable to remove temporary uploaded CV file')


def validate_saved_cv(filepath):
    """Validate file content after saving it, then clean up invalid uploads."""
    try:
        is_valid = validate_cv_file(str(filepath))
    except OSError:
        app.logger.exception('Unable to validate uploaded CV file')
        is_valid = False

    if not is_valid:
        remove_upload_file(filepath)
    return is_valid

def embedding_to_bytes(embedding: np.ndarray) -> bytes:

    return embedding.tobytes()

def bytes_to_embedding(data: bytes) -> np.ndarray:

    return np.frombuffer(data, dtype=np.float32)

@app.route('/')
def index():

    db = get_db()
    cv_count = db.execute("SELECT COUNT(*) FROM cvs").fetchone()[0]
    job_count = db.execute("SELECT COUNT(*) FROM jobs").fetchone()[0]
    match_count = db.execute("SELECT COUNT(DISTINCT job_id) FROM matches").fetchone()[0]

    recent_cvs = db.execute(
        "SELECT id, filename, extracted_skills, uploaded_at FROM cvs ORDER BY uploaded_at DESC LIMIT 5"
    ).fetchall()

    recent_jobs = db.execute(
        "SELECT id, title, required_skills, created_at FROM jobs ORDER BY created_at DESC LIMIT 5"
    ).fetchall()

    return render_template('index.html',
                           cv_count=cv_count,
                           job_count=job_count,
                           match_count=match_count,
                           recent_cvs=recent_cvs,
                           recent_jobs=recent_jobs)

@app.route('/upload-cv', methods=['GET', 'POST'])
def upload_cv():

    if request.method == 'POST':
        if 'files' not in request.files:
            flash('Dosya seçilmedi.', 'error')
            return redirect(request.url)

        files = request.files.getlist('files')
        uploaded_count = 0
        errors = []

        for file in files:
            filepath = None
            if file.filename == '':
                continue

            if not allowed_file(file.filename):
                errors.append(f"'{file.filename}' — sadece PDF ve DOCX dosyaları desteklenmektedir.")
                continue

            try:

                filename = secure_filename(file.filename)
                unique_name = f"{uuid.uuid4().hex[:8]}_{filename}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
                file.save(filepath)

                if not validate_saved_cv(Path(filepath)):
                    errors.append(f"'{filename}' — dosya içeriği geçersiz veya desteklenmiyor.")
                    continue

                parse_result = parse_file(filepath)
                raw_text = parse_result['text']
                warnings = parse_result['warnings']

                if not raw_text or len(raw_text.strip()) < 50:
                    errors.append(f"'{filename}' — PDF/DOCX'den yeterli metin çıkarılamadı.")
                    remove_upload_file(Path(filepath))
                    continue

                if warnings:
                    session['upload_warnings'] = warnings

                extracted_info = nlp_engine.extract_cv_info(raw_text)

                cleaned = clean_text(raw_text)

                embedding = nlp_engine.encode_text(raw_text)

                cv_id = str(uuid.uuid4())
                db = get_db()
                db.execute(
                    """INSERT INTO cvs (id, filename, file_path, original_text, cleaned_text,
                       extracted_skills, timeline, experience_months, skill_recency, metadata, embedding, uploaded_at)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (cv_id, filename, unique_name, raw_text, cleaned,
                     json.dumps(extracted_info.get('skills', [])),
                     json.dumps(extracted_info.get('timeline', [])),
                     extracted_info.get('total_experience_months', 0),
                     json.dumps(extracted_info.get('skill_recency', {})),
                     json.dumps({'warnings': warnings}),
                     embedding_to_bytes(embedding),
                     datetime.now().isoformat())
                )
                db.commit()

                uploaded_count += 1

            except Exception:
                app.logger.exception('CV upload failed')
                remove_upload_file(Path(filepath) if filepath else None)
                errors.append(f"'{file.filename}' — CV işlenemedi. Lütfen geçerli bir PDF veya DOCX deneyin.")

        if uploaded_count > 0:
            flash(f'{uploaded_count} CV başarıyla yüklendi ve analiz edildi!', 'success')
        if errors:
            for err in errors:
                flash(err, 'error')

        return redirect(url_for('upload_cv'))

    db = get_db()
    cvs = db.execute(
        "SELECT id, filename, file_path, extracted_skills, uploaded_at, experience_months FROM cvs ORDER BY uploaded_at DESC"
    ).fetchall()

    return render_template('upload_cv.html', cvs=cvs)

@app.route('/job-analysis', methods=['GET', 'POST'])
def job_analysis():

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        custom_skills = request.form.get('custom_skills', '').strip()
        linkedin_url = request.form.get('linkedin_url', '').strip()

        if linkedin_url:
            if not validate_linkedin_url(linkedin_url):
                flash('Geçersiz LinkedIn URL formatı.', 'error')
                return redirect(request.url)

            scrape_res = scrape_linkedin_job(linkedin_url)
            if not scrape_res['success']:
                flash('LinkedIn ilanı çekilemedi: ' + scrape_res.get('error', 'Hata'), 'error')
                return redirect(request.url)

            title = scrape_res['title']
            description = scrape_res['description']

            custom_skills = ""

        if not title:
            flash('İş ilanı başlığı boş olamaz.', 'error')
            return redirect(request.url)

        if not description:
            flash('İş ilanı açıklaması boş olamaz.', 'error')
            return redirect(request.url)

        try:

            cleaned_desc = clean_text(description)

            extracted_skills = nlp_engine.extract_skills(description)

            if custom_skills:
                user_skills = [s.strip().lower() for s in custom_skills.split(',') if s.strip()]
                extracted_skills = sorted(list(set(extracted_skills + user_skills)))

            job_reqs = nlp_engine.extract_job_requirements(description)
            must_have = job_reqs['must_have_skills']
            nice_to_have = job_reqs['nice_to_have_skills']

            if custom_skills:
                for s in user_skills:
                    if s not in must_have:
                        must_have.append(s)
                must_have = sorted(list(set(must_have)))

            embedding = nlp_engine.encode_text(description)

            job_id = str(uuid.uuid4())
            db = get_db()
            db.execute(
                """INSERT INTO jobs (id, title, description, cleaned_description,
                   required_skills, must_have_skills, nice_to_have_skills,
                   embedding, created_at)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (job_id, title, description, cleaned_desc,
                 json.dumps(extracted_skills),
                 json.dumps(must_have),
                 json.dumps(nice_to_have),
                 embedding_to_bytes(embedding),
                 datetime.now().isoformat())
            )
            db.commit()

            must_count = len(must_have)
            nice_count = len(nice_to_have)
            total = len(extracted_skills)
            flash(
                f'İş ilanı "{title}" kaydedildi! '
                f'{total} yetenek tespit edildi '
                f'({must_count} zorunlu, {nice_count} tercih edilen).',
                'success'
            )
            return redirect(url_for('job_analysis'))

        except Exception as e:
            flash(f'Hata: {str(e)}', 'error')
            return redirect(request.url)

    db = get_db()
    jobs = db.execute(
        "SELECT id, title, required_skills, must_have_skills, description, created_at FROM jobs ORDER BY created_at DESC"
    ).fetchall()

    return render_template('job_analysis.html', jobs=jobs)

@app.route('/match/<job_id>', methods=['POST'])
def run_match(job_id):

    db = get_db()

    job = db.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        flash('İş ilanı bulunamadı.', 'error')
        return redirect(url_for('job_analysis'))

    cvs = db.execute("SELECT * FROM cvs").fetchall()
    if not cvs:
        flash('Henüz yüklenmiş CV bulunmuyor. Önce CV yükleyin.', 'error')
        return redirect(url_for('upload_cv'))

    db.execute("DELETE FROM matches WHERE job_id = ?", (job_id,))

    for cv_row in cvs:

        cv = dict(cv_row)

        cv_metadata = json.loads(cv.get('metadata') or '{}')
        pdf_warnings = cv_metadata.get('warnings', [])

        ats_result = nlp_engine.calculate_ats_score(
            cv_text=cv.get('original_text', ''),
            job_description=job['description'],
            pdf_warnings=pdf_warnings
        )

        job_skills = nlp_engine.extract_skills(job['description'])
        cv_skills = json.loads(cv.get('extracted_skills', '[]'))
        gap = nlp_engine.semantic_gap_analysis(job_skills, cv_skills)

        match_id = str(uuid.uuid4())
        db.execute(
            """INSERT INTO matches (
                id, job_id, cv_id, match_score, matching_skills,
                semantic_matches, missing_skills, extra_skills,
                timeline_gaps, experience_score, coverage_percent,
                text_similarity, format_score, keyword_score,
                section_score, language_match, missing_must_haves,
                sections_found, cv_lang, job_lang, title_match_bonus,
                is_disqualified, penalty_applied, is_pretty_resume, detail_metrics
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                match_id, job_id, cv['id'], ats_result['final_score'],
                json.dumps(gap['exact_matches']),
                json.dumps(gap['semantic_matches']),
                json.dumps(gap['missing_skills']),
                json.dumps(gap['extra_skills']),
                json.dumps(cv.get('gaps_detected', [])),
                0,
                gap['coverage_percent'],
                0,
                ats_result['format_score'],
                ats_result['keyword_score'],
                ats_result['section_score'],
                1 if ats_result['language_match'] else 0,
                json.dumps(ats_result['missing_must_haves']),
                json.dumps(ats_result['sections_found']),
                ats_result['cv_lang'],
                ats_result['job_lang'],
                ats_result['title_match_bonus'],
                ats_result['is_disqualified'],
                ats_result['penalty_applied'],
                ats_result['is_pretty_resume'],
                json.dumps(ats_result['detail_metrics'])
            )
        )

    db.commit()
    return redirect(url_for('results', job_id=job_id))

@app.route('/results/<job_id>')
def results(job_id):

    db = get_db()

    job = db.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not job:
        flash('İş ilanı bulunamadı.', 'error')
        return redirect(url_for('index'))

    matches = db.execute("""
        SELECT m.*, c.filename, c.extracted_skills as cv_skills, c.experience_months
        FROM matches m
        JOIN cvs c ON m.cv_id = c.id
        WHERE m.job_id = ?
        ORDER BY m.is_disqualified ASC, m.match_score DESC
    """, (job_id,)).fetchall()

    return render_template('results.html', job=job, matches=matches)

@app.route('/api/delete-cv/<cv_id>', methods=['DELETE'])
def delete_cv(cv_id):

    db = get_db()
    cv = db.execute("SELECT file_path FROM cvs WHERE id = ?", (cv_id,)).fetchone()
    if not cv:
        return jsonify({'success': False, 'message': 'CV bulunamadı.'}), 404

    filepath = get_upload_path(cv['file_path'])
    staged_path = None
    try:
        db.execute('BEGIN')
        db.execute("DELETE FROM matches WHERE cv_id = ?", (cv_id,))
        db.execute("DELETE FROM cvs WHERE id = ?", (cv_id,))

        if filepath and filepath.is_file():
            staged_path = filepath.with_name(f'.deleting-{uuid.uuid4().hex}-{filepath.name}')
            filepath.replace(staged_path)

        db.commit()
    except (OSError, sqlite3.Error):
        db.rollback()
        if staged_path and staged_path.exists():
            try:
                staged_path.replace(filepath)
            except OSError:
                app.logger.exception('Unable to restore CV file after database deletion failure')
        app.logger.exception('CV deletion failed')
        return jsonify({'success': False, 'message': 'CV silinemedi. Lütfen tekrar deneyin.'}), 500

    if staged_path:
        try:
            staged_path.unlink()
        except OSError:
            app.logger.exception('CV database records deleted but staged file cleanup failed')
            return jsonify({
                'success': True,
                'message': 'CV silindi; dosya temizliği başarısız oldu. Sistem yöneticisine başvurun.',
                'cleanup_pending': True,
            })

    return jsonify({'success': True, 'message': 'CV silindi.'})

@app.route('/api/delete-job/<job_id>', methods=['DELETE'])
def delete_job(job_id):

    db = get_db()
    db.execute("DELETE FROM matches WHERE job_id = ?", (job_id,))
    db.execute("DELETE FROM jobs WHERE id = ?", (job_id,))
    db.commit()
    return jsonify({'success': True, 'message': 'İş ilanı silindi.'})

@app.route('/download_cv/<cv_id>')
def download_cv_file(cv_id):

    db = get_db()
    cv = db.execute("SELECT file_path, filename FROM cvs WHERE id = ?", (cv_id,)).fetchone()
    if not cv or not cv['file_path']:
        flash('Dosya bulunamadı veya eski bir CV olduğu için silinmiş.', 'error')
        return redirect(url_for('upload_cv'))

    filepath = get_upload_path(cv['file_path'])
    if not filepath or not filepath.is_file():
        flash('Fiziksel dosya sunucuda bulunamadı.', 'error')
        return redirect(url_for('upload_cv'))

    from flask import send_from_directory
    return send_from_directory(filepath.parent, filepath.name, as_attachment=True, download_name=cv['filename'])

@app.route('/api/job-search/delete-session/<session_id>', methods=['DELETE'])
def job_search_delete_session(session_id):

    db = get_db()
    db.execute("DELETE FROM job_search_sessions WHERE id = ?", (session_id,))
    db.commit()
    return jsonify({'success': True, 'message': 'Oturum silindi.'})

@app.route('/api/stats')
def api_stats():

    db = get_db()
    return jsonify({
        'cv_count': db.execute("SELECT COUNT(*) FROM cvs").fetchone()[0],
        'job_count': db.execute("SELECT COUNT(*) FROM jobs").fetchone()[0],
        'match_count': db.execute("SELECT COUNT(DISTINCT job_id) FROM matches").fetchone()[0],
    })

@app.route('/job-search')
def job_search():

    db = get_db()

    profiles = db.execute(
        "SELECT id, original_filename, extracted_skills, created_at FROM user_profiles ORDER BY created_at DESC"
    ).fetchall()

    sessions = db.execute("""
        SELECT s.*, p.original_filename
        FROM job_search_sessions s
        JOIN user_profiles p ON s.profile_id = p.id
        ORDER BY s.created_at DESC LIMIT 10
    """).fetchall()
    return render_template('job_search.html', profiles=profiles, sessions=sessions)

@app.route('/api/job-search/upload-cv', methods=['POST'])
def job_search_upload_cv():

    filepath = None

    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'Dosya seçilmedi.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'Dosya seçilmedi.'}), 400

    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'Sadece PDF ve DOCX dosyaları desteklenmektedir.'}), 400

    try:

        filename = secure_filename(file.filename)
        unique_name = f"{uuid.uuid4().hex[:8]}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
        file.save(filepath)

        if not validate_saved_cv(Path(filepath)):
            return jsonify({'success': False, 'error': 'Dosya içeriği geçersiz veya desteklenmiyor.'}), 400

        parse_result = parse_file(filepath)
        raw_text = parse_result['text']
        warnings = parse_result['warnings']
        if not raw_text or len(raw_text.strip()) < 50:
            remove_upload_file(Path(filepath))
            return jsonify({'success': False, 'error': 'PDF/DOCX\'den yeterli metin çıkarılamadı.'}), 400

        profile_data = extractive_cv_gen.parse_cv(raw_text)

        if not profile_data:
            remove_upload_file(Path(filepath))
            return jsonify({'success': False, 'error': 'Profil parse edilemedi.'}), 500

        skills = nlp_engine.extract_skills(raw_text)

        profile_id = str(uuid.uuid4())
        db = get_db()
        db.execute(
            """INSERT INTO user_profiles (id, original_filename, original_text,
               profile_data, extracted_skills, created_at)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (profile_id, filename, raw_text,
             json.dumps(profile_data, ensure_ascii=False),
             json.dumps(skills),
             datetime.now().isoformat())
        )
        db.commit()

        remove_upload_file(Path(filepath))

        return {
            'success': True,
            'profile_id': profile_id,
            'profile': profile_data,
            'skills': skills,
            'filename': filename,
            'warnings': warnings
        }

    except Exception:
        app.logger.exception('Job search CV upload failed')
        remove_upload_file(Path(filepath) if filepath else None)
        return jsonify({'success': False, 'error': 'CV işlenemedi. Lütfen tekrar deneyin.'}), 500

@app.route('/api/job-search/parse-job', methods=['POST'])
def job_search_parse_job():

    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'error': 'Veri gönderilmedi.'}), 400

    url = data.get('url', '').strip()
    manual_text = data.get('manual_text', '').strip()
    manual_title = data.get('manual_title', '').strip()
    manual_company = data.get('manual_company', '').strip()

    if manual_text:

        result = parse_job_text(manual_text, manual_title, manual_company)
    elif url:

        if not validate_linkedin_url(url):
            return jsonify({
                'success': False,
                'error': 'Geçersiz LinkedIn URL formatı. Lütfen geçerli bir LinkedIn iş ilanı linki girin.'
            }), 400
        result = scrape_linkedin_job(url)
    else:
        return jsonify({'success': False, 'error': 'LinkedIn URL veya ilan metni girilmelidir.'}), 400

    if result['success'] and result.get('description'):
        extracted_skills = nlp_engine.extract_skills(result['description'])
        result['extracted_skills'] = extracted_skills

    return jsonify(result)

@app.route('/api/job-search/generate-cv', methods=['POST'])
def job_search_generate_cv():

    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'error': 'Veri gönderilmedi.'}), 400

    profile_id = data.get('profile_id', '').strip()
    job_data = data.get('job_data')

    if not profile_id:
        return jsonify({'success': False, 'error': 'Profil ID gerekli.'}), 400
    if not job_data:
        return jsonify({'success': False, 'error': 'İlan bilgileri gerekli.'}), 400

    db = get_db()
    profile = db.execute("SELECT * FROM user_profiles WHERE id = ?", (profile_id,)).fetchone()
    if not profile:
        return jsonify({'success': False, 'error': 'Profil bulunamadı.'}), 404

    profile_data = json.loads(profile['profile_data'])

    try:
        optimized_cv_data = extractive_cv_gen.generate_tailored_cv(profile_data, job_data)
    except Exception as e:
        return jsonify({'success': False, 'error': f'CV oluşturulurken hata: {str(e)}'}), 500

    session_id = str(uuid.uuid4())
    db.execute(
        """INSERT INTO job_search_sessions (id, profile_id, job_url, job_data,
           optimized_cv, status, created_at)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (session_id, profile_id,
         job_data.get('url', ''),
         json.dumps(job_data, ensure_ascii=False),
         json.dumps(optimized_cv_data, ensure_ascii=False),
         'completed',
         datetime.now().isoformat())
    )
    db.commit()

    return jsonify({
        'success': True,
        'session_id': session_id,
        'cv_data': optimized_cv_data
    })

def generate_docx_buffer(cv_data):

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv_data.get('full_name', 'İsimsiz').upper())
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact = cv_data.get('contact', {})
    contact_parts = []
    if contact.get('email'): contact_parts.append(contact['email'])
    if contact.get('phone'): contact_parts.append(contact['phone'])
    if contact.get('location'): contact_parts.append(contact['location'])
    if contact.get('linkedin'): contact_parts.append(contact['linkedin'])

    if contact_parts:
        contact_p = doc.add_paragraph(' • '.join(contact_parts))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = cv_data.get('headers', {})
    summary = cv_data.get('professional_summary', '')
    if summary:
        doc.add_heading(headers.get('summary', 'SUMMARY'), level=1)
        doc.add_paragraph(summary)

    experiences = cv_data.get('experience', [])
    if experiences:
        doc.add_heading(headers.get('experience', 'EXPERIENCE'), level=1)
        for exp in experiences:
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp.get('title', '')}")
            title_run.bold = True
            p.add_run(f" — {exp.get('company', '')}")

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(2)
            meta_text = []
            if exp.get('location'): meta_text.append(exp['location'])
            if exp.get('period'): meta_text.append(exp['period'])
            meta_p.add_run(' | '.join(meta_text)).italic = True

            for ach in exp.get('achievements', []):
                bullet = doc.add_paragraph(ach, style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.25)

    education = cv_data.get('education', [])
    if education:
        doc.add_heading(headers.get('education', 'EDUCATION'), level=1)
        for edu in education:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')}").bold = True
            p.add_run(f" | {edu.get('institution', '')}")
            if edu.get('period'):
                p.add_run(f" ({edu.get('period', '')})").italic = True
            if edu.get('details'):
                doc.add_paragraph(edu['details'])

    skills = cv_data.get('skills', {})
    if skills:
        doc.add_heading(headers.get('skills', 'SKILLS'), level=1)
        skill_groups = skills.get('groups', [])
        if skill_groups:
            for group in skill_groups:
                doc.add_paragraph(", ".join(group))
        else:
            if skills.get('primary'):
                p = doc.add_paragraph()
                p.add_run("Temel: ").bold = True
                p.add_run(", ".join(skills['primary']))
            if skills.get('secondary'):
                p = doc.add_paragraph()
                p.add_run("Ek: ").bold = True
                p.add_run(", ".join(skills['secondary']))
            if skills.get('tools'):
                p = doc.add_paragraph()
                p.add_run("Araçlar: ").bold = True
                p.add_run(", ".join(skills['tools']))

    if cv_data.get('languages'):
        doc.add_heading(headers.get('languages', 'LANGUAGES'), level=1)
        langs = [f"{l.get('language','')} ({l.get('level','')})" for l in cv_data['languages']]
        p = doc.add_paragraph()
        p.add_run(", ".join(langs))

    if cv_data.get('certifications'):
        doc.add_heading(headers.get('certifications', 'CERTIFICATIONS'), level=1)
        for cert in cv_data.get('certifications', []):
            cert_text = f"{cert.get('name','')} - {cert.get('issuer','')}"
            if cert.get('year'): cert_text += f" ({cert['year']})"
            doc.add_paragraph(cert_text, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/api/job-search/download-cv/<session_id>')
def job_search_download_cv(session_id):

    file_format = request.args.get('format', 'pdf').lower()

    db = get_db()
    session = db.execute("SELECT * FROM job_search_sessions WHERE id = ?", (session_id,)).fetchone()
    if not session:
        return jsonify({'success': False, 'error': 'Oturum bulunamadı.'}), 404

    cv_data = json.loads(session['optimized_cv'])
    name = cv_data.get('full_name', 'cv_ats_optimized')
    safe_name = name.replace(' ', '_').lower()[:30]

    if file_format == 'docx':
        buffer = generate_docx_buffer(cv_data)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{safe_name}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    db = get_db()
    session = db.execute("SELECT * FROM job_search_sessions WHERE id = ?", (session_id,)).fetchone()
    if not session:
        return jsonify({'success': False, 'error': 'Oturum bulunamadı.'}), 404

    cv_data = json.loads(session['optimized_cv'])

    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=15*mm, bottomMargin=15*mm
    )

    font_name = 'Helvetica'
    try:
        import os
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        arial_path = r'C:\Windows\Fonts\arial.ttf'
        arial_bold_path = r'C:\Windows\Fonts\arialbd.ttf'
        arial_italic_path = r'C:\Windows\Fonts\ariali.ttf'

        if os.path.exists(arial_path):
            pdfmetrics.registerFont(TTFont('Arial', arial_path))
            pdfmetrics.registerFont(TTFont('Arial-Bold', arial_bold_path))
            pdfmetrics.registerFont(TTFont('Arial-Italic', arial_italic_path))
            font_name = 'Arial'
    except:
        pass

    font_bold = f"{font_name}-Bold" if font_name == 'Arial' else 'Helvetica-Bold'
    font_italic = f"{font_name}-Italic" if font_name == 'Arial' else 'Helvetica-Oblique'

    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        name='CVName', fontSize=18, fontName=font_bold,
        spaceAfter=6, textColor=HexColor('#000000'), alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        name='CVContact', fontSize=10, fontName=font_name,
        spaceAfter=10, textColor=HexColor('#333333'), alignment=TA_CENTER
    ))
    styles.add(ParagraphStyle(
        name='CVSection', fontSize=12, fontName=font_bold,
        spaceBefore=12, spaceAfter=4, textColor=HexColor('#000000'),
        borderPadding=(0, 0, 2, 0)
    ))
    styles.add(ParagraphStyle(
        name='CVSubtitle', fontSize=11, fontName=font_bold,
        spaceAfter=2, textColor=HexColor('#000000')
    ))
    styles.add(ParagraphStyle(
        name='CVMeta', fontSize=9, fontName=font_italic,
        spaceAfter=3, textColor=HexColor('#444444')
    ))
    styles.add(ParagraphStyle(
        name='CVBody', fontSize=10.5, fontName=font_name,
        spaceAfter=4, textColor=HexColor('#000000'), leading=14,
        alignment=TA_LEFT
    ))
    styles.add(ParagraphStyle(
        name='CVBullet', fontSize=10.5, fontName=font_name,
        spaceAfter=3, textColor=HexColor('#000000'), leading=13,
        leftIndent=15, bulletIndent=5
    ))

    elements = []

    name = cv_data.get('full_name', 'İsimsiz')
    elements.append(Paragraph(name, styles['CVName']))

    contact = cv_data.get('contact', {})
    contact_parts = []
    if contact.get('email'):
        contact_parts.append(contact['email'])
    if contact.get('phone'):
        contact_parts.append(contact['phone'])
    if contact.get('location'):
        contact_parts.append(contact['location'])
    if contact.get('linkedin'):
        contact_parts.append(contact['linkedin'])
    if contact_parts:
        elements.append(Paragraph(' • '.join(contact_parts), styles['CVContact']))

    elements.append(HRFlowable(width='100%', thickness=1, color=HexColor('#c0c0e0'), spaceAfter=8))

    headers = cv_data.get('headers', {})
    summary = cv_data.get('professional_summary', '')
    if summary:
        elements.append(Paragraph(headers.get('summary', 'PROFESSIONAL SUMMARY'), styles['CVSection']))
        elements.append(Paragraph(summary, styles['CVBody']))

    experiences = cv_data.get('experience', [])
    if experiences:
        elements.append(Paragraph(headers.get('experience', 'EXPERIENCE'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for exp in experiences:
            title_company = f"<b>{exp.get('title', '')}</b> — {exp.get('company', '')}"
            elements.append(Paragraph(title_company, styles['CVSubtitle']))
            meta_parts = []
            if exp.get('location'):
                meta_parts.append(exp['location'])
            if exp.get('period'):
                meta_parts.append(exp['period'])
            if meta_parts:
                elements.append(Paragraph(' | '.join(meta_parts), styles['CVMeta']))
            for ach in exp.get('achievements', []):
                elements.append(Paragraph(f'• {ach}', styles['CVBullet']))
            elements.append(Spacer(1, 4))

    education = cv_data.get('education', [])
    if education:
        elements.append(Paragraph(headers.get('education', 'EDUCATION'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for edu in education:
            elements.append(Paragraph(f"<b>{edu.get('degree', '')}</b>", styles['CVSubtitle']))
            meta_parts = []
            if edu.get('institution'):
                meta_parts.append(edu['institution'])
            if edu.get('period'):
                meta_parts.append(edu['period'])
            if meta_parts:
                elements.append(Paragraph(' | '.join(meta_parts), styles['CVMeta']))
            if edu.get('details'):
                elements.append(Paragraph(edu['details'], styles['CVBody']))
            elements.append(Spacer(1, 3))

    skills = cv_data.get('skills', {})
    if skills:
        elements.append(Paragraph(headers.get('skills', 'SKILLS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))

        skill_groups = skills.get('groups', [])
        if skill_groups:
            for group in skill_groups:
                elements.append(Paragraph(', '.join(group), styles['CVBody']))
        else:

            if skills.get('primary'):
                elements.append(Paragraph(f"<b>Temel:</b> {', '.join(skills['primary'])}", styles['CVBody']))
            if skills.get('secondary'):
                elements.append(Paragraph(f"<b>Ek:</b> {', '.join(skills['secondary'])}", styles['CVBody']))
            if skills.get('tools'):
                elements.append(Paragraph(f"<b>Araçlar:</b> {', '.join(skills['tools'])}", styles['CVBody']))

    languages = cv_data.get('languages', [])
    if languages:
        elements.append(Paragraph(headers.get('languages', 'LANGUAGES'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        lang_text = ', '.join([f"{l.get('language', '')} ({l.get('level', '')})" for l in languages])
        elements.append(Paragraph(lang_text, styles['CVBody']))

    certs = cv_data.get('certifications', [])
    if certs:
        elements.append(Paragraph(headers.get('certifications', 'CERTIFICATIONS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for cert in certs:
            cert_text = cert.get('name', '')
            if cert.get('issuer'):
                cert_text += f" — {cert['issuer']}"
            if cert.get('year'):
                cert_text += f" ({cert['year']})"
            elements.append(Paragraph(f'• {cert_text}', styles['CVBullet']))

    projects = cv_data.get('projects', [])
    if projects:
        elements.append(Paragraph(headers.get('projects', 'PROJECTS'), styles['CVSection']))
        elements.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#d0d0e8'), spaceAfter=6))
        for proj in projects:
            elements.append(Paragraph(f"<b>{proj.get('name', '')}</b>", styles['CVSubtitle']))
            if proj.get('description'):
                elements.append(Paragraph(proj['description'], styles['CVBody']))
            if proj.get('technologies'):
                elements.append(Paragraph(f"Teknolojiler: {', '.join(proj['technologies'])}", styles['CVMeta']))
            elements.append(Spacer(1, 3))

    doc.build(elements)
    buffer.seek(0)

    safe_name = name.replace(' ', '_').lower()[:30]
    pdf_filename = f"cv_{safe_name}_{session_id[:8]}.pdf"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=pdf_filename,
        mimetype='application/pdf'
    )

@app.route('/api/job-search/delete-profile/<profile_id>', methods=['DELETE'])
def delete_profile(profile_id):

    db = get_db()
    db.execute("DELETE FROM job_search_sessions WHERE profile_id = ?", (profile_id,))
    db.execute("DELETE FROM user_profiles WHERE id = ?", (profile_id,))
    db.commit()
    return jsonify({'success': True, 'message': 'Profil silindi.'})

@app.template_filter('parse_json')
def parse_json_filter(value):

    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return []

@app.template_filter('format_date')
def format_date_filter(value):

    try:
        dt = datetime.fromisoformat(value)
        return dt.strftime('%d.%m.%Y %H:%M')
    except (ValueError, TypeError):
        return value

if __name__ == '__main__':
    # NOT: Veritabanı tablolarının kurulduğundan emin olup, yerel sunucuyu 5000 portunda başlatıyoruz.
    with app.app_context():
        init_db()
    print("\n" + "=" * 50)
    print("  HireLens - Semantic Talent Matcher & ATS Engine")
    print("  Server running at: http://localhost:5000")
    print("=" * 50 + "\n")
    app.run(port=5000)
